"""Tests for references module, learn_from_reference skill, and /teach /files commands."""

from __future__ import annotations

from pathlib import Path
from unittest.mock import patch

import pytest

from halia.references import (
    delete_reference,
    get_reference_path,
    list_ref_files,
    search_ref_files,
    store_reference,
    store_url_reference,
)
from halia.skills.reference import LearnFromReference


def _make_file(tmp_path: Path, name: str, content: str = "test content") -> Path:
    f = tmp_path / name
    f.write_text(content)
    return f


def _db(tmp_path: Path) -> Path:
    """A temp DB path for isolated tests."""
    return tmp_path / "test.db"


# ── store_reference ───────────────────────────────────────────────────────────


def test_store_reference_basic(tmp_path: Path) -> None:
    f = _make_file(tmp_path, "template.md", "# Test Flow\nStep 1, Step 2")
    ref = store_reference(str(f), db_path=_db(tmp_path))
    assert ref.filename == "template.md"
    assert ref.file_type == ".md"
    assert ref.profile == ""
    assert ref.size_bytes > 0
    assert ref.id


def test_store_reference_with_profile(tmp_path: Path) -> None:
    f = _make_file(tmp_path, "qa-flow.pdf", "PDF content here")
    ref = store_reference(
        str(f), profile="qa", description="Test flow format", db_path=_db(tmp_path)
    )
    assert ref.profile == "qa"
    assert ref.description == "Test flow format"
    assert ref.file_type == ".pdf"


def test_store_reference_unsupported_format(tmp_path: Path) -> None:
    f = _make_file(tmp_path, "image.exe", "binary")
    with pytest.raises(ValueError, match="unsupported format"):
        store_reference(str(f), db_path=_db(tmp_path))


def test_store_reference_missing_file(tmp_path: Path) -> None:
    with pytest.raises(FileNotFoundError):
        store_reference("/nonexistent/file.txt", db_path=_db(tmp_path))


def test_store_reference_too_large(tmp_path: Path) -> None:
    f = tmp_path / "huge.txt"
    f.write_text("x" * (6 * 1024 * 1024))  # 6MB
    with pytest.raises(ValueError, match="too large"):
        store_reference(str(f), db_path=_db(tmp_path))


# ── list_ref_files ────────────────────────────────────────────────────────────


def test_list_ref_files_empty(tmp_path: Path) -> None:
    assert list_ref_files(db_path=_db(tmp_path)) == []


def test_list_ref_files_with_data(tmp_path: Path) -> None:
    _make_file(tmp_path, "a.md")
    _make_file(tmp_path, "b.pdf")
    db = _db(tmp_path)
    store_reference(str(tmp_path / "a.md"), profile="qa", db_path=db)
    store_reference(str(tmp_path / "b.pdf"), db_path=db)
    refs = list_ref_files(db_path=db)
    assert len(refs) == 2


def test_list_ref_files_filter_by_profile(tmp_path: Path) -> None:
    _make_file(tmp_path, "a.md")
    _make_file(tmp_path, "b.pdf")
    db = _db(tmp_path)
    store_reference(str(tmp_path / "a.md"), profile="qa", db_path=db)
    store_reference(str(tmp_path / "b.pdf"), profile="finance", db_path=db)
    refs = list_ref_files(profile="qa", db_path=db)
    assert len(refs) == 1
    assert refs[0].filename == "a.md"


# ── search_ref_files ──────────────────────────────────────────────────────────


def test_search_ref_files(tmp_path: Path) -> None:
    _make_file(tmp_path, "login-flow.md")
    _make_file(tmp_path, "email-template.txt")
    db = _db(tmp_path)
    store_reference(str(tmp_path / "login-flow.md"), description="Login test flow", db_path=db)
    store_reference(str(tmp_path / "email-template.txt"), db_path=db)
    refs = search_ref_files("login", db_path=db)
    assert len(refs) == 1
    assert refs[0].filename == "login-flow.md"


def test_search_ref_files_no_match(tmp_path: Path) -> None:
    _make_file(tmp_path, "test.md")
    db = _db(tmp_path)
    store_reference(str(tmp_path / "test.md"), db_path=db)
    assert search_ref_files("nonexistent", db_path=db) == []


# ── get_reference_path ────────────────────────────────────────────────────────


def test_get_reference_path(tmp_path: Path) -> None:
    f = _make_file(tmp_path, "data.txt", "hello")
    db = _db(tmp_path)
    ref = store_reference(str(f), db_path=db)
    path = get_reference_path(ref.id, db_path=db)
    assert path is not None
    assert path.exists()
    assert path.read_text() == "hello"


def test_get_reference_path_not_found(tmp_path: Path) -> None:
    assert get_reference_path("nonexistent", db_path=_db(tmp_path)) is None


# ── delete_reference ──────────────────────────────────────────────────────────


def test_delete_reference(tmp_path: Path) -> None:
    f = _make_file(tmp_path, "to_delete.txt")
    db = _db(tmp_path)
    ref = store_reference(str(f), db_path=db)
    assert delete_reference(ref.id, db_path=db) is True
    assert get_reference_path(ref.id, db_path=db) is None
    assert list_ref_files(db_path=db) == []


def test_delete_reference_not_found(tmp_path: Path) -> None:
    assert delete_reference("nonexistent", db_path=_db(tmp_path)) is False


def test_delete_reference_removes_stored_file(tmp_path: Path) -> None:
    # Regression: delete must unlink the physical file under ~/.halia/files/, not just
    # the DB row. (It used to look up the original basename, never the stored name.)
    f = _make_file(tmp_path, "phys.txt", "unique-content-for-delete-regression")
    db = _db(tmp_path)
    ref = store_reference(str(f), db_path=db)
    stored = get_reference_path(ref.id, db_path=db)
    assert stored is not None and stored.exists()
    assert delete_reference(ref.id, db_path=db) is True
    assert not stored.exists()  # physical file gone


# ── learn_from_reference skill ────────────────────────────────────────────────


def test_learn_from_reference_no_files(tmp_path: Path) -> None:
    skill = LearnFromReference()
    with patch("halia.references.list_ref_files", return_value=[]):
        result = skill.run({})
    assert "no reference files" in result


def test_learn_from_reference_with_files(tmp_path: Path) -> None:
    f = _make_file(tmp_path, "template.md", "# Format\nStep 1, Step 2")
    db = _db(tmp_path)
    ref = store_reference(str(f), db_path=db)
    with patch("halia.references.get_reference_path", return_value=f):
        with patch("halia.references.list_ref_files", return_value=[ref]):
            skill = LearnFromReference()
            result = skill.run({})
    assert "template.md" in result
    assert "# Format" in result
    assert "Step 1" in result


def test_learn_from_reference_filter_by_profile(tmp_path: Path) -> None:
    db = _db(tmp_path)
    ref_qa = store_reference(
        str(_make_file(tmp_path, "qa-flow.md", "QA format")),
        profile="qa", db_path=db,
    )
    store_reference(
        str(_make_file(tmp_path, "finance-flow.md", "Finance format")),
        profile="finance", db_path=db,
    )
    qa_file = tmp_path / "qa-flow.md"
    with patch("halia.references.list_ref_files", return_value=[ref_qa]), \
         patch("halia.references.get_reference_path", return_value=qa_file):
        skill = LearnFromReference()
        result = skill.run({"profile": "qa"})
    assert "qa-flow.md" in result
    assert "finance-flow.md" not in result


def test_learn_from_reference_specific_file(tmp_path: Path) -> None:
    db = _db(tmp_path)
    ref_a = store_reference(str(_make_file(tmp_path, "a.md", "Content A")), db_path=db)
    ref_b = store_reference(str(_make_file(tmp_path, "b.md", "Content B")), db_path=db)
    file_a = tmp_path / "a.md"
    file_b = tmp_path / "b.md"
    with patch("halia.references.list_ref_files", return_value=[ref_a, ref_b]), \
         patch("halia.references.get_reference_path",
               side_effect=lambda rid, **kw: file_a if rid == ref_a.id else file_b):
        skill = LearnFromReference()
        result = skill.run({"filename": "a.md"})
    assert "Content A" in result
    assert "Content B" not in result


def test_learn_from_reference_extracts_docx(tmp_path: Path) -> None:
    # Regression: a taught .docx must return real extracted text, not binary mojibake.
    from docx import Document

    doc = Document()
    doc.add_paragraph("UNIQUE_DOCX_MARKER heading")
    doc.add_paragraph("second line body")
    docx_path = tmp_path / "template.docx"
    doc.save(str(docx_path))

    ref = store_reference(str(docx_path), profile="qa", db_path=_db(tmp_path))
    with patch("halia.references.list_ref_files", return_value=[ref]), \
         patch("halia.references.get_reference_path", return_value=docx_path):
        result = LearnFromReference().run({})
    assert "UNIQUE_DOCX_MARKER" in result
    assert "second line body" in result


def test_learn_from_reference_extracts_xlsx(tmp_path: Path) -> None:
    # Regression: a taught .xlsx must surface real cell values, not binary mojibake.
    import openpyxl

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.append(["Name", "Score"])
    ws.append(["Alice", 95])
    xlsx_path = tmp_path / "data.xlsx"
    wb.save(str(xlsx_path))

    ref = store_reference(str(xlsx_path), db_path=_db(tmp_path))
    with patch("halia.references.list_ref_files", return_value=[ref]), \
         patch("halia.references.get_reference_path", return_value=xlsx_path):
        result = LearnFromReference().run({})
    assert "Name" in result and "Score" in result
    assert "Alice" in result


# ── store_url_reference (teaching a URL) ────────────────────────────────────────


def test_store_url_reference(tmp_path: Path) -> None:
    db = _db(tmp_path)
    ref = store_url_reference(
        "https://example.com/reporting-standard",
        profile="finance",
        description="reporting standard",
        db_path=db,
        fetcher=lambda _u: "Rule 1: disclose everything",
    )
    assert ref.url == "https://example.com/reporting-standard"
    assert ref.profile == "finance"
    assert ref.file_type == ".md"
    assert ref.filename.startswith("example.com")
    stored = get_reference_path(ref.id, db_path=db)
    assert stored is not None and "Rule 1: disclose everything" in stored.read_text()


def test_store_url_reference_empty_raises(tmp_path: Path) -> None:
    with pytest.raises(ValueError, match="no readable text"):
        store_url_reference("https://example.com", db_path=_db(tmp_path), fetcher=lambda _u: "   ")


def test_url_reference_listed_with_url(tmp_path: Path) -> None:
    db = _db(tmp_path)
    store_url_reference("https://example.com/a", db_path=db, fetcher=lambda _u: "content")
    refs = list_ref_files(db_path=db)
    assert len(refs) == 1
    assert refs[0].url == "https://example.com/a"


def test_learn_from_reference_shows_url_source(tmp_path: Path) -> None:
    db = _db(tmp_path)
    ref = store_url_reference(
        "https://example.com/guide", profile="qa", db_path=db,
        fetcher=lambda _u: "Best practice: write clear steps",
    )
    stored = get_reference_path(ref.id, db_path=db)
    with patch("halia.references.list_ref_files", return_value=[ref]), \
         patch("halia.references.get_reference_path", return_value=stored):
        result = LearnFromReference().run({})
    assert "source: https://example.com/guide" in result
    assert "Best practice" in result


def test_skill_registered() -> None:
    from halia.skills import available_skills
    assert "learn_from_reference" in available_skills()
