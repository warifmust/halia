"""Tests for make_docx (markdown -> Word)."""

from typing import Any

from halia.skills.export import MakeDocx, render_markdown_docx

_MD = """# Class Report

A short **summary**.

## Scores

| Student | Average |
| --- | --- |
| Aisha | 78.3 |
| Chong | 90.0 |

- First point.
- Second point.
"""


def _all_text(doc: Any) -> str:
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)
    return " ".join(parts)


def test_render_has_headings_list_and_table() -> None:
    doc = render_markdown_docx(_MD)
    text = _all_text(doc)
    assert "Class Report" in text
    assert "First point." in text
    assert "Chong" in text and "90.0" in text
    assert len(doc.tables) == 1


def test_bold_becomes_a_bold_run() -> None:
    doc = render_markdown_docx("A **bold** word.")
    runs = [r for p in doc.paragraphs for r in p.runs]
    assert any(r.text == "bold" and r.bold for r in runs)


def test_make_docx_writes_doc_and_source(tmp_path: Any) -> None:
    out = tmp_path / "report.docx"
    result = MakeDocx().run({"path": str(out), "content": _MD})
    assert "wrote Word document" in result
    assert out.read_bytes().startswith(b"PK")  # docx is a zip
    assert (tmp_path / "report.md").read_text() == _MD


def test_make_docx_supports_unicode(tmp_path: Any) -> None:
    out = tmp_path / "u.docx"
    r = MakeDocx().run({"path": str(out), "content": "# Ringkasan\n\nKehadiran ↑ 好"})
    assert "wrote" in r
    from docx import Document

    assert "Ringkasan" in " ".join(p.text for p in Document(str(out)).paragraphs)


def test_make_docx_validates(tmp_path: Any) -> None:
    assert "content" in MakeDocx().run({"path": str(tmp_path / "x.docx"), "content": " "})


def test_make_docx_honors_permission_floor(tmp_path: Any) -> None:
    blocked = tmp_path / "id_rsa.docx"
    r = MakeDocx().run({"path": str(blocked), "content": "# Secret\n\nhi"})
    assert r.startswith("blocked:")
    assert not blocked.exists()


def test_make_docx_is_dangerous_and_wired() -> None:
    from halia.presets import get_preset
    from halia.skills import available_skills

    assert MakeDocx().dangerous is True
    assert "make_docx" in available_skills()
    for vertical in ("finance", "research", "education"):
        assert "make_docx" in get_preset(vertical).skills
