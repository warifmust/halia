"""Compliance hook — check_requirements.

A requirement-coverage check: given a document and a list of required terms/clauses,
report which are present (with a citable snippet) and which are MISSING. It reads the
FILE itself — not text relayed by the model — so the check runs against ground truth,
not a paraphrase. That's the deterministic part.

Honest boundary: presence is deterministic; whether a *found* clause is ADEQUATE is a
judgment the model must flag as such — "the policy mentions data retention" is not the
same as "the retention clause meets the standard".
"""

from __future__ import annotations

from pathlib import Path
from typing import Any

from halia.permissions.guard import check_readable

_SNIPPET_PAD = 55


def _extract_text(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        from pypdf import PdfReader

        return "\n".join(page.extract_text() or "" for page in PdfReader(str(path)).pages)
    if suffix == ".docx":
        from docx import Document

        doc = Document(str(path))
        parts = [p.text for p in doc.paragraphs]
        for table in doc.tables:
            for row in table.rows:
                parts.append(" | ".join(cell.text for cell in row.cells))
        return "\n".join(parts)
    return path.read_text(encoding="utf-8", errors="replace")


class CheckRequirements:
    name = "check_requirements"
    description = (
        "Check a document for required terms/clauses (coverage check). Given a file `path` "
        "and a list of `requirements`, reports each as FOUND (with a surrounding snippet to "
        "cite) or MISSING. Reads the file directly, so the check is against the real "
        "document. Presence is deterministic; whether a found clause is ADEQUATE is a "
        "separate judgment you must state as such."
    )
    dangerous = False
    untrusted = False
    parameters: dict[str, Any] = {
        "type": "object",
        "additionalProperties": False,
        "properties": {
            "path": {"type": "string", "description": "Path to the document (.pdf/.docx/text)."},
            "requirements": {
                "type": "array",
                "items": {"type": "string"},
                "description": "Required terms or clauses to look for.",
            },
        },
        "required": ["path", "requirements"],
    }

    def run(self, args: dict[str, Any]) -> str:
        raw = args.get("path")
        requirements = args.get("requirements")
        if not isinstance(raw, str) or not raw.strip():
            return "error: 'path' is required"
        if not isinstance(requirements, list) or not requirements:
            return "error: 'requirements' must be a non-empty array of strings"
        reqs = [r for r in requirements if isinstance(r, str) and r.strip()]
        if not reqs:
            return "error: 'requirements' must contain at least one non-empty string"

        path = Path(raw).expanduser()
        check_readable(path)
        if not path.is_file():
            return f"error: not a file: {path}"
        try:
            text = _extract_text(path)
        except (OSError, ValueError) as exc:
            return f"error reading {path}: {exc}"

        haystack = text.lower()
        found: list[str] = []
        missing: list[str] = []
        for req in reqs:
            idx = haystack.find(req.strip().lower())
            if idx >= 0:
                start = max(0, idx - _SNIPPET_PAD)
                end = idx + len(req) + _SNIPPET_PAD
                snippet = " ".join(text[start:end].split())
                found.append(f'- "{req}" → …{snippet}…')
            else:
                missing.append(f'- "{req}"')

        lines = [f"Coverage: {len(found)}/{len(reqs)} requirements found in {path.name}."]
        if found:
            lines.append("\nFOUND (presence only — review adequacy):")
            lines.extend(found)
        if missing:
            lines.append(f"\nMISSING ({len(missing)}):")
            lines.extend(missing)
        else:
            lines.append("\nMISSING: none.")
        return "\n".join(lines)
