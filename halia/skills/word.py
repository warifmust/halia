"""Word ingestion — read_docx.

Extracts text (paragraphs + tables) from a .docx via python-docx (already a
dependency for make_docx). Closes the read/write asymmetry: halia could write
Word but not read it — and contracts/policies live in Word.
"""

from __future__ import annotations

from pathlib import Path
from typing import Any

from halia.permissions.guard import check_readable

_MAX_CHARS = 10_000


class ReadDocx:
    name = "read_docx"
    description = (
        "Read a Word .docx file and return its text (paragraphs and tables, truncated)."
    )
    dangerous = False
    parameters: dict[str, Any] = {
        "type": "object",
        "additionalProperties": False,
        "properties": {
            "path": {"type": "string", "description": "Path to the .docx file."},
            "max_chars": {
                "type": "integer",
                "description": "Max characters of text to return (default 10000).",
            },
        },
        "required": ["path"],
    }

    def run(self, args: dict[str, Any]) -> str:
        raw = args.get("path")
        if not isinstance(raw, str) or not raw.strip():
            return "error: 'path' is required and must be a non-empty string"
        path = Path(raw).expanduser()
        check_readable(path)
        if not path.is_file():
            return f"error: not a file: {path}"

        max_chars = args.get("max_chars", _MAX_CHARS)
        if not isinstance(max_chars, int) or max_chars <= 0:
            max_chars = _MAX_CHARS

        from docx import Document
        from docx.opc.exceptions import PackageNotFoundError

        try:
            doc = Document(str(path))
        except (PackageNotFoundError, OSError, ValueError) as exc:
            return f"error reading .docx: {exc}"

        parts: list[str] = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                if any(cells):
                    parts.append(" | ".join(cells))

        combined = "\n".join(parts)
        if not combined.strip():
            return "no text found in the document."
        if len(combined) > max_chars:
            combined = combined[:max_chars] + "\n… [truncated]"
        return f"extracted text (paragraphs + tables):\n\n{combined}"
